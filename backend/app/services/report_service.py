"""全流程追溯报告生成: 聚合 15 项 -> Jinja2 HTML -> PDF(weasyprint 优先, xhtml2pdf 降级)。

PDF 三级降级策略:
  1. weasyprint (高质量 CSS 渲染, 需系统库 pango/cairo)
  2. xhtml2pdf + reportlab CID 字体 (无需系统库, 纯 Python)
  3. 纯 HTML (以上均不可用时, 保留完整内容)
报告入 report_records + file_objects, 带版本与数据快照。
"""
from __future__ import annotations

import os
import statistics
from io import BytesIO
from datetime import datetime, timezone

from sqlalchemy.orm import Session

from app.models import (
    AuditLog, DiagnosisResult, EvaluationResult, FactorDictionary, FileObject,
    ImportBatch, Measurement, MLModel, Recommendation, ReportRecord, SamplingPoint,
    RemediationCase, Site, StandardThreshold, TechnologyLibrary, ThresholdRule,
    WorkflowAttachment, WorkflowRecord,
)
from app.services.file_service import save_bytes
from app.services.workflow_service import STAGE_NAME, get_stages

from app.core.config import resource_root

ROOT = resource_root()
TEMPLATE_DIR = os.path.join(ROOT, "reporting", "templates")
TEMPLATE_VERSION = "tpl_v0.1"


def _factor_summary(db: Session, site_id: int) -> list[dict]:
    rows = (db.query(FactorDictionary.factor_code, FactorDictionary.level1_category,
                     Measurement.value, Measurement.unit)
            .join(Measurement, Measurement.factor_id == FactorDictionary.id)
            .filter(Measurement.site_id == site_id).all())
    agg: dict[str, dict] = {}
    for code, cat, val, unit in rows:
        if val is None:
            continue
        d = agg.setdefault(code, {"factor": code, "category": cat, "vals": [], "unit": unit})
        d["vals"].append(val)
    out = []
    for d in agg.values():
        vs = d["vals"]
        out.append({"factor": d["factor"], "category": d["category"], "count": len(vs),
                    "min": round(min(vs), 3), "mean": round(statistics.mean(vs), 3),
                    "max": round(max(vs), 3), "unit": d["unit"]})
    out.sort(key=lambda x: x["factor"])
    return out


def _coverage_summary(db: Session, site_id: int, n_points: int) -> dict:
    factor_count = (db.query(FactorDictionary.factor_code)
                    .join(Measurement, Measurement.factor_id == FactorDictionary.id)
                    .filter(Measurement.site_id == site_id)
                    .distinct().count())
    n_meas = db.query(Measurement).filter_by(site_id=site_id).count()
    denominator = n_points * factor_count if n_points and factor_count else 0
    coverage = round(n_meas / denominator * 100, 2) if denominator else 0
    return {
        "factor_count": factor_count,
        "observed_cells": n_meas,
        "expected_cells": denominator,
        "coverage_pct": coverage,
        "missing_pct": round(100 - coverage, 2) if denominator else 0,
        "note": "measurements 长表只保存实测项; 未检测字段不会伪造为 0。",
    }


def _standard_versions(db: Session) -> list[dict]:
    rows = (db.query(StandardThreshold.standard_code, StandardThreshold.standard_name,
                     StandardThreshold.version, StandardThreshold.source_reference)
            .distinct().all())
    return [{"standard_code": c, "standard_name": n, "version": v,
             "source_reference": r} for c, n, v, r in rows]


def _remediation_cases(db: Session, site: Site, limit: int = 8) -> list[dict]:
    q = db.query(RemediationCase)
    if site.pollution_type:
        token = {"heavy_metal": "HM", "organic": "OP", "composite": "HM+OP"}.get(
            site.pollution_type, site.pollution_type)
        q = q.filter(RemediationCase.pollution_type.contains(token))
    rows = q.limit(limit).all()
    if not rows:
        rows = db.query(RemediationCase).limit(limit).all()
    return [{
        "case_id": r.case_id,
        "site_type": r.site_type,
        "region": r.region,
        "land_use": r.land_use,
        "pollution_type": r.pollution_type,
        "pollutants": r.pollutants,
        "remediation_technology": r.remediation_technology,
        "technology_category": r.technology_category,
        "cost_level": r.cost_level,
        "effectiveness": r.effectiveness,
        "limitation": r.limitation,
        "secondary_risk": r.secondary_risk,
        "evidence_source": r.evidence_source,
        "doi": r.doi,
    } for r in rows]


def docx_emu_width(doc) -> int:
    """返回 DOCX 正文可用宽度(EMU 单位), 用于图片自适应页宽。"""
    try:
        section = doc.sections[0]
        # 页宽 - 左右边距
        usable = section.page_width - section.left_margin - section.right_margin
        return int(usable) if usable and usable > 0 else 5500000  # 兜底约 15cm
    except Exception:  # noqa: BLE001
        return 5500000


def _embed_docx_image(doc, data_url: str | None, caption: str) -> None:
    """把 base64 PNG 嵌入 DOCX; data_url 为空/损坏时静默跳过( P1-1 DOCX 同步 PDF 图件)。"""
    if not data_url or not data_url.startswith("data:image/png;base64,"):
        return
    import base64 as _b64
    from io import BytesIO as _BIO
    try:
        img_bytes = _b64.b64decode(data_url.split(",", 1)[1])
        doc.add_picture(_BIO(img_bytes), width=docx_emu_width(doc))
        p = doc.add_paragraph(caption)
        p.italic = True
    except Exception:  # noqa: BLE001
        doc.add_paragraph(f"[{caption} 渲染失败]")


def _render_points_map_png(coord_points: list, exceed_by_point: dict[int, float],
                           exceed_factor: dict[int, str] | None = None) -> str | None:
    """v0.2: 8级色阶采样点风险散点图(与前端/API一致), 返回 base64 PNG。"""
    if not coord_points:
        return None
    try:
        # 优先用 static_map_renderer
        from app.services.static_map_renderer import (
            _exc_color, COLOR_8_LEVELS, _exc_label, _get_cjk_font,
        )
    except ImportError:
        # fallback: 内联渲染
        _exc_color = None

    try:
        import base64 as _b64
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    font = _get_cjk_font() if _get_cjk_font else None
    if font:
        plt.rcParams["font.family"] = font

    lons = [float(p.longitude) for p in coord_points]
    lats = [float(p.latitude) for p in coord_points]
    vals = [exceed_by_point.get(p.id, 0.0) for p in coord_points]
    factors = (exceed_factor or {})

    # v0.2: 8级色阶
    if _exc_color:
        colors = [_exc_color(v) for v in vals]
    else:
        def _fallback_color(v):
            if v < 1: return "#16a34a"
            if v < 3: return "#facc15"
            if v < 10: return "#f59e0b"
            if v < 30: return "#ea580c"
            if v < 80: return "#dc2626"
            if v < 200: return "#9f1239"
            return "#6b0f1a"
        colors = [_fallback_color(v) for v in vals]
    sizes = [24 if v >= 1 else 14 for v in vals]
    fig, ax = plt.subplots(figsize=(7.5, 5), dpi=120)
    ax.scatter(lons, lats, c=colors, s=sizes, alpha=0.85, edgecolors="white", linewidths=0.5, zorder=5)

    # 图例 8级
    from matplotlib.patches import Patch as _Patch
    _leg = []
    if _exc_label:
        for th, clr in COLOR_8_LEVELS:
            _leg.append(_Patch(facecolor=clr, edgecolor="white", label=_exc_label(float(th))))
    else:
        _leg = [
            _Patch(facecolor="#16a34a", edgecolor="white", label="未超标"),
            _Patch(facecolor="#facc15", edgecolor="white", label="轻度 1-3x"),
            _Patch(facecolor="#f59e0b", edgecolor="white", label="中度 3-10x"),
            _Patch(facecolor="#ea580c", edgecolor="white", label="偏重 10-30x"),
            _Patch(facecolor="#dc2626", edgecolor="white", label="重度 30-80x"),
            _Patch(facecolor="#9f1239", edgecolor="white", label="极重 80-200x"),
            _Patch(facecolor="#6b0f1a", edgecolor="white", label="超极重 >200x"),
            _Patch(facecolor="#64748b", edgecolor="white", label="无数据"),
        ]
    ax.legend(handles=_leg, loc="lower right", fontsize=6, ncol=2, framealpha=0.9)

    if len(set(lons)) > 1:
        ax.set_xlim(min(lons) - (max(lons)-min(lons))*0.08, max(lons) + (max(lons)-min(lons))*0.08)
    if len(set(lats)) > 1:
        ax.set_ylim(min(lats) - (max(lats)-min(lats))*0.08, max(lats) + (max(lats)-min(lats))*0.08)
    ax.set_xlabel("经度" if font else "Longitude", fontsize=8)
    ax.set_ylabel("纬度" if font else "Latitude", fontsize=8)
    ax.set_title(f"采样点超标风险分布（{len(coord_points)} 点位, 8级色阶）" if font
                 else f"Exceedance risk ({len(coord_points)} pts, 8-level)", fontsize=10, fontweight="bold")
    ax.tick_params(labelsize=7)
    ax.grid(True, linestyle="--", alpha=0.3)
    from datetime import datetime as _dt, timezone as _tz
    wm = f"8级色阶 | 渲染: {_dt.now(_tz.utc).strftime('%Y-%m-%d %H:%M UTC')} | 底图: 无(离线坐标散点)"
    fig.text(0.5, 0.01, wm, ha="center", fontsize=5.5, color="#888")
    fig.tight_layout(rect=[0, 0.03, 1, 0.97])
    from io import BytesIO
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=120, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return f"data:image/png;base64,{_b64.b64encode(buf.read()).decode('ascii')}"


def _render_shap_figure_png(top_factors: list, site_name: str) -> str | None:
    """用 matplotlib 画模型贡献份额横向条形图(科研配图风格)。

    ``contribution_scope=local_point`` 时明确标注为真实点位局部解释；
    否则标注为训练集全局背景贡献。
    正向(加重)=npg红 #E64B35, 负向(缓解)=npg蓝 #4DBBD5; 去顶右边框, 数值标注。"""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:
        return None
    from matplotlib import font_manager as _fm
    chinese_fonts = ["Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "WenQuanYi Zen Hei"]
    available_fonts = {font.name for font in _fm.fontManager.ttflist}
    selected_font = next((font for font in chinese_fonts if font in available_fonts), None)
    if selected_font:
        plt.rcParams["font.sans-serif"] = [selected_font]
        plt.rcParams["axes.unicode_minus"] = False
    if not top_factors:
        return None
    import io, base64
    facts = top_factors[:8][::-1]  # 横向 bar 倒序(最重要在顶部)
    names = [str(f.get("factor", "?")) for f in facts]
    vals = [float(f.get("importance", 0) or 0) for f in facts]
    dirs = [str(f.get("direction", "")) for f in facts]
    is_local = any(f.get("contribution_scope") == "local_point" for f in facts)
    scope_label = (
        "真实点位局部 SHAP 贡献份额" if is_local else "训练集全局背景贡献份额"
    ) if selected_font else ("Local point SHAP share" if is_local else "Global background share")
    colors = ["#E64B35" if d == "positive" else "#4DBBD5" for d in dirs]
    fig, ax = plt.subplots(figsize=(7, 3.4), dpi=150)
    ax.barh(range(len(names)), vals, color=colors, height=0.62, edgecolor="white", linewidth=0.6)
    ax.set_yticks(range(len(names)))
    ax.set_yticklabels(names, fontsize=9)
    ax.set_xlabel(scope_label, fontsize=9)
    title = (f"关键障碍因子{scope_label} — {site_name}" if selected_font
             else "Barrier factor contribution")
    ax.set_title(title, fontsize=10.5, pad=8, color="#222")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.tick_params(length=0)
    for i, v in enumerate(vals):
        ax.text(v, i, f" {v:.3f}", va="center", fontsize=7.5, color="#555")
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode("ascii")


def _render_eda_figure_png(factor_summary: list) -> str | None:
    """各因子浓度均值与最大值对比柱状图(matplotlib), 嵌 DOCX 检测数据摘要章节。"""
    if not factor_summary:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return None
    from matplotlib import font_manager as _fm
    _cn_fonts = ["PingFang SC", "Heiti SC", "STHeiti", "Arial Unicode MS",
                 "Noto Sans CJK SC", "WenQuanYi Zen Hei", "SimHei", "Microsoft YaHei"]
    _avail = {_f.name for _f in _fm.fontManager.ttflist}
    _has_cn = any(_f in _avail for _f in _cn_fonts)
    if _has_cn:
        plt.rcParams["font.sans-serif"] = _cn_fonts
        plt.rcParams["axes.unicode_minus"] = False
    facts = factor_summary[:12]
    names = [f["factor"] for f in facts]
    means = [float(f.get("mean") or 0) for f in facts]
    maxs = [float(f.get("max") or 0) for f in facts]
    x = list(range(len(names)))
    w = 0.38
    fig, ax = plt.subplots(figsize=(7, 3.4), dpi=120)
    ax.bar([i - w / 2 for i in x], means, w, label="均值", color="#3680ae")
    ax.bar([i + w / 2 for i in x], maxs, w, label="最大值", color="#e98184")
    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=40, ha="right", fontsize=8)
    ax.set_ylabel("浓度", fontsize=9)
    ax.set_title("各因子浓度均值与最大值对比(EDA)", fontsize=10, color="#222")
    ax.legend(fontsize=8)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    import io as _io, base64 as _b64
    buf = _io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return "data:image/png;base64," + _b64.b64encode(buf.getvalue()).decode("ascii")


def collect(db: Session, site_id: int, version: str) -> dict:
    site = db.get(Site, site_id)
    if site is None:
        raise ValueError(f"场地不存在: {site_id}")

    batch = (db.query(ImportBatch).filter_by(site_id=site_id)
             .order_by(ImportBatch.id.desc()).first())
    points = db.query(SamplingPoint).filter_by(site_id=site_id).all()
    n_meas = db.query(Measurement).filter_by(site_id=site_id).count()
    vr = (batch.validation_report or {}) if batch else {}

    diag = (db.query(DiagnosisResult).filter_by(site_id=site_id)
            .order_by(DiagnosisResult.id.desc()).first())
    diag_ctx = None
    if diag:
        from app.models import DiagnosisFactorDetail
        model = db.get(MLModel, diag.model_id) if diag.model_id else None
        details = (db.query(DiagnosisFactorDetail, FactorDictionary)
                   .join(FactorDictionary, DiagnosisFactorDetail.factor_id == FactorDictionary.id)
                   .filter(DiagnosisFactorDetail.diagnosis_id == diag.id,
                           DiagnosisFactorDetail.sampling_point_id.is_(None))
                   .order_by(DiagnosisFactorDetail.rank).all())
        # Round9 P0-3.4: 区分 KOS / RF+SHAP, 不再默认 model_name="RF"
        diag_method = diag.diagnosis_method or "rf_shap"
        if diag_method == "kos":
            model_name_default = "KOS-P3Alpha"
        else:
            model_name_default = "RF"
        diag_ctx = {
            "model_name": model.model_name if model else model_name_default,
            "model_version": model.version if model else (diag.model_version or "—"),
            "metrics": (model.metrics if model else {}) or {},
            "data_version": diag.data_version, "summary": diag.summary,
            "top_factors": [{"rank": d.rank, "factor": fd.factor_name,
                             "category": fd.level1_category,
                             "importance": d.importance, "direction": d.direction}
                            for d, fd in details],
            "method": diag_method,
            "track": diag.track,
            "subset": diag.subset,
        }
        # Round9 P0-3.4: KOS 记录从 result_payload 读完整审计信息(五分量/模型贡献度/开放集)
        if diag_method == "kos" and diag.result_payload:
            rp = diag.result_payload
            diag_ctx["top_factors"] = [
                {
                    "rank": item.get("rank"),
                    "factor": item.get("factor"),
                    "category": item.get("category") or "—",
                    "kos_score": item.get("KOS"),
                    "direction": item.get("direction") or "—",
                    "components": item.get("components") or {},
                    "value": item.get("value"),
                    "threshold_value": item.get("threshold_value"),
                    "threshold_standard": item.get("threshold_standard") or "",
                    "max_exceedance_ratio": (
                        (rp.get("per_point_stats") or {})
                        .get(item.get("factor"), {})
                        .get("max_exceedance_ratio")
                    ),
                }
                for item in (rp.get("key_obstacles") or [])
            ]
            diag_ctx["kos"] = {
                "key_obstacles": rp.get("key_obstacles", []),
                "model_contribution": rp.get("model_contribution", []),
                "model_contribution_scope": rp.get("model_contribution_scope", "global_model"),
                "local_shap_status": rp.get("local_shap_status", ""),
                "decision_point_id": rp.get("decision_point_id"),
                "decision_point_code": rp.get("decision_point_code"),
                "decision_point_selection": rp.get("decision_point_selection"),
                "per_point_stats": rp.get("per_point_stats", {}),
                "n_sampling_points": rp.get("n_sampling_points", 0),
                "open_set": rp.get("open_set", {}),
                "open_set_summary": rp.get("open_set_summary", {}),
                "family_warnings": rp.get("family_warnings", []),
                "unknown_alerts": rp.get("unknown_alerts", []),
                "model_attention_factors": rp.get("model_attention_factors", []),
                "kos_weights": rp.get("kos_weights", {}),
                "interpretation_note": rp.get("interpretation_note", ""),
                "threshold_version": rp.get("threshold_version", ""),
                "review_required": rp.get("review_required", False),
                "data_quality_flags": rp.get("data_quality_flags", []),
                "coverage": rp.get("coverage", 0),
                "limitations": rp.get("limitations", ""),
            }
            # Round9 P0-3.4: 开放集四层(替代硬编码空 [])
            os_data = rp.get("open_set", {}) or {}
            formal_obstacles = [k.get("factor") for k in rp.get("key_obstacles", [])
                                 if k.get("threshold_resolution_status") == "resolved"]
        else:
            formal_obstacles = []
        # 占位会被后续覆盖

    evals = {e.eval_type: e for e in
             db.query(EvaluationResult).filter_by(site_id=site_id)
             .order_by(EvaluationResult.id.desc()).all()[::-1]}
    recon = []
    for et, title in (("reconstruction_prod", "生产功能重构"),
                      ("reconstruction_eco", "生态功能重构")):
        e = evals.get(et)
        if e:
            recon.append({"title": title, "score": e.score, "grade": e.grade,
                          "limiting_factors": e.limiting_factors or [],
                          "explanation": e.explanation})
    ssui_e = evals.get("ssui")
    ssui_ctx = ({"score": ssui_e.score, "grade": ssui_e.grade,
                 "explanation": ssui_e.explanation} if ssui_e else None)

    recs = (db.query(Recommendation).filter_by(site_id=site_id)
            .order_by(Recommendation.rank).all())
    tech = {t.id: t for t in db.query(TechnologyLibrary).all()}
    rec_ctx = []
    for r in recs:
        t = tech.get(r.technology_id)
        rec_ctx.append({
            "rank": r.rank,
            "technology": t.tech_name if t else "—",
            "match_score": r.match_score,
            "reason": r.reason or "",
            "cost_level": t.cost_level if t else None,
            "duration_level": t.duration_level if t else None,
            "limitations": t.limitations if t else None,
            "forbidden_conditions": t.forbidden_conditions if t else None,
            "secondary_risk": t.secondary_risk if t else None,
        })

    workflow = get_stages(db, site_id)
    attachments = []
    for w in db.query(WorkflowRecord).filter_by(site_id=site_id).all():
        for a in db.query(WorkflowAttachment).filter_by(workflow_record_id=w.id).all():
            fo = db.get(FileObject, a.file_object_id)
            attachments.append({"stage_name": STAGE_NAME.get(w.stage), "file_role": a.file_role,
                                "original_name": fo.original_name if fo else "—"})

    logs = (db.query(AuditLog).order_by(AuditLog.id.desc()).limit(10).all())
    audit_ctx = [{"created_at": str(a.created_at), "action": a.action,
                  "resource_type": a.resource_type, "resource_id": a.resource_id,
                  "result": a.result} for a in logs]

    coord_points = [p for p in points if p.longitude is not None and p.latitude is not None]
    if coord_points:
        lons = [float(p.longitude) for p in coord_points]
        lats = [float(p.latitude) for p in coord_points]
        bounds = {"min_lon": round(min(lons), 6), "max_lon": round(max(lons), 6),
                  "min_lat": round(min(lats), 6), "max_lat": round(max(lats), 6)}
    else:
        bounds = None

    # v0.2: 每个采样点的最大超标倍数 + 最严重因子
    exceed_by_point: dict[int, float] = {}
    exceed_factor: dict[int, str] = {}  # pid → factor_code
    if coord_points:
        th_rows = (db.query(Measurement.sampling_point_id, Measurement.value,
                            ThresholdRule.threshold_max, FactorDictionary.factor_code)
                   .join(FactorDictionary, Measurement.factor_id == FactorDictionary.id)
                   .join(ThresholdRule, ThresholdRule.factor_id == FactorDictionary.id)
                   .filter(Measurement.site_id == site_id,
                           ThresholdRule.threshold_max != None,
                           ThresholdRule.threshold_max > 0,
                           Measurement.sampling_point_id != None).all())
        for pid, val, tmax, fcode in th_rows:
            if val is None:
                continue
            ratio = float(val) / float(tmax)
            if ratio > exceed_by_point.get(pid, 0.0):
                exceed_by_point[pid] = ratio
                exceed_factor[pid] = fcode  # v0.2: 保留最严重因子
                exceed_by_point[pid] = ratio
    map_image = _render_points_map_png(coord_points, exceed_by_point, exceed_factor)
    contribution_rows = (diag_ctx or {}).get("top_factors", [])
    if diag_ctx and diag_ctx.get("method") == "kos":
        contribution_rows = [
            {
                "factor": item.get("factor"),
                "importance": item.get("contribution"),
                "direction": item.get("direction"),
                "contribution_scope": item.get("contribution_scope"),
            }
            for item in (diag_ctx.get("kos", {}).get("model_contribution") or [])
        ]
    shap_image = _render_shap_figure_png(contribution_rows, site.name)

    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    return {
        "site": {"site_code": site.site_code, "name": site.name,
                 "pollution_type": site.pollution_type, "land_use_type": site.land_use_type,
                 "province": site.province, "city": site.city,
                 "longitude": float(site.longitude) if site.longitude is not None else None,
                 "latitude": float(site.latitude) if site.latitude is not None else None},
        "data_source": {"source_file": batch.source_file if batch else None,
                        "batch_id": batch.id if batch else None,
                        "status": batch.status if batch else None,
                        "script_version": batch.script_version if batch else None,
                        "n_points": len(points), "n_measurements": n_meas},
        "sampling_points": [{"point_code": p.point_code, "region": p.region,
                             "longitude": float(p.longitude) if p.longitude is not None else None,
                             "latitude": float(p.latitude) if p.latitude is not None else None,
                             "depth_top_cm": p.depth_top_cm, "depth_bottom_cm": p.depth_bottom_cm,
                             "soil_type": p.soil_type} for p in points],
        "factor_summary": _factor_summary(db, site_id),
        "map_summary": {"n_points": len(points), "n_coord_points": len(coord_points),
                        "coverage_pct": round(len(coord_points) / max(len(points), 1) * 100, 2),
                        "bounds": bounds,
                        "map_image": map_image,
                        "shap_image": shap_image,
                        "note": "交互式地图由系统图层接口生成, 按污染物筛选并以超标倍数分级。上方静态图件由 matplotlib 离线渲染(无瓦片底图)。"},
        "coverage": _coverage_summary(db, site_id, len(points)),
        "validation": {"passed": vr.get("passed", True), "n_errors": vr.get("n_errors", 0),
                       "n_warnings": vr.get("n_warnings", 0), "n_exceed": vr.get("n_exceed", 0),
                       "exceed_factors": (vr.get("summary", {}) or {}).get("exceed_factors", [])},
        "diagnosis": diag_ctx, "reconstruction": recon, "ssui": ssui_ctx,
        "recommendations": rec_ctx,
        "remediation_cases": _remediation_cases(db, site),
        "standard_versions": _standard_versions(db),
        "workflow": workflow,
        "attachments": attachments, "audit_logs": audit_ctx,
        # Round9 P0-3.4: 开放集四层 — 从 KOS result_payload 注入(替代硬编码空 [])
        # KOS 记录: 优先从 result_payload.open_set 读四层
        # 非 KOS 记录: 保留空(模板渲染为"无")
        "formal_obstacles": (
            (diag_ctx or {}).get("kos", {}).get("open_set", {}).get("formal_obstacles", [])
            if (diag_ctx and diag_ctx.get("method") == "kos") else []),
        "model_candidates": (
            (diag_ctx or {}).get("kos", {}).get("open_set", {}).get("model_candidates", [])
            if (diag_ctx and diag_ctx.get("method") == "kos") else []),
        "family_alerts": (
            (diag_ctx or {}).get("kos", {}).get("family_warnings", [])
            if (diag_ctx and diag_ctx.get("method") == "kos") else []),
        "unknown_measured": (
            (diag_ctx or {}).get("kos", {}).get("unknown_alerts", [])
            if (diag_ctx and diag_ctx.get("method") == "kos") else []),
        "report": {"version": version, "template_version": TEMPLATE_VERSION,
                   "data_version": diag.data_version if diag else f"site{site_id}",
                   "standard_version": "GB15618/GB36600/HJ25.5-2018",
                   "generated_at": now},
    }


def render_html(context: dict) -> str:
    from jinja2 import Environment, FileSystemLoader, select_autoescape
    env = Environment(loader=FileSystemLoader(TEMPLATE_DIR),
                      autoescape=select_autoescape(["html"]))
    return env.get_template("traceability_report.html").render(**context)


def _plain_reportlab_pdf(html: str) -> bytes | None:
    """无 Cairo 环境的最终 PDF 降级路径。

    只使用 ReportLab 的 PDF canvas 与内置中文 CID 字体，避免 Windows
    演示机因缺少 libcairo 导致报告接口直接 500。完整富文本仍由前两级渲染器负责。
    """
    try:
        import html as html_lib
        import re
        from io import BytesIO
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas

        text = re.sub(r"<(script|style)\b[^>]*>.*?</\1>", "", html,
                      flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(r"</?(?:h[1-6]|p|div|tr|li|br|section|table|ul|ol)[^>]*>",
                      "\n", text, flags=re.IGNORECASE)
        text = html_lib.unescape(re.sub(r"<[^>]+>", "", text))
        lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
        lines = [line for line in lines if line]

        buf = BytesIO()
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        c = canvas.Canvas(buf, pagesize=A4)
        width, height = A4
        y = height - 42
        c.setFont("STSong-Light", 9)
        for line in lines:
            # CID 字体下按字符数折行，避免依赖系统字体测量或 Cairo。
            for start in range(0, len(line), 52):
                if y < 42:
                    c.showPage()
                    c.setFont("STSong-Light", 9)
                    y = height - 42
                c.drawString(36, y, line[start:start + 52])
                y -= 13
        c.save()
        return buf.getvalue()
    except Exception:
        return None


def html_to_pdf(html: str) -> bytes | None:
    """HTML → PDF: WeasyPrint → xhtml2pdf → 纯 ReportLab。"""
    # 第一级: weasyprint (CSS3 完整支持, 中文排版好)
    try:
        from weasyprint import HTML
        return HTML(string=html).write_pdf()
    except (ImportError, OSError):
        pass
    # 第二级: xhtml2pdf + reportlab UnicodeCID 字体 (纯 Python, 无需系统库)
    try:
        from io import BytesIO
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from xhtml2pdf import pisa
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        except Exception:
            pass
        buf = BytesIO()
        result = pisa.CreatePDF(src=html, dest=buf, encoding="utf-8")
        if result.err:
            return None
        return buf.getvalue()
    except (ImportError, OSError):
        pass
    return _plain_reportlab_pdf(html)


# ── Round10: 专业 DOCX 样式常量 ─────────────────────────────────
# 颜色方案: 深蓝主色调, 与系统 UI 一致 (#0f3d6e)
_HEADER_BG = "0F3D6E"      # 表头深蓝底
_HEADER_FG = "FFFFFF"       # 表头白字
_ROW_ALT = "F5F7FA"         # 交替行底色
_BORDER = "B8C4D0"          # 表格边框
_ACCENT_RED = "B91C1C"      # 强调红
_TITLE_FONT = "SimHei"       # 标题字体（黑体）
_BODY_FONT = "SimSun"        # 正文字体（宋体）
_WATERMARK_TEXT = "SRS 监管系统"


def _set_cell_shading(cell, color: str):
    """设置单元格底色（python-docx shading）。"""
    from docx.oxml.ns import qn
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def _style_table(table, header_rows: int = 1):
    """对已填充数据的表格应用专业样式: 表头深蓝底白字 + 交替行底色。"""
    from docx.oxml.ns import qn
    from docx.shared import Pt
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            # 边框
            tcPr = cell._element.get_or_add_tcPr()
            borders = tcPr.makeelement(qn("w:tcBorders"), {})
            for edge in ("top", "left", "bottom", "right"):
                el = borders.makeelement(qn(f"w:{edge}"), {
                    qn("w:val"): "single",
                    qn("w:sz"): "4",
                    qn("w:color"): _BORDER,
                })
                borders.append(el)
            tcPr.append(borders)
            # 字体
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = _BODY_FONT
            if i < header_rows:
                _set_cell_shading(cell, _HEADER_BG)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = None  # reset
                        from docx.shared import RGBColor
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
            elif i % 2 == 0:
                _set_cell_shading(cell, _ROW_ALT)


def _make_kv_table(doc, rows_data: list[tuple[str, object]], col_widths=(0.28, 0.72)):
    """创建键值对表格并应用样式。返回 table 对象。"""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.autofit = True
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = str(k)
        table.rows[i].cells[1].text = "" if v is None else str(v)
    _style_table(table, header_rows=0)
    # 对键列应用浅灰底色
    for row in table.rows:
        _set_cell_shading(row.cells[0], "F0F4F8")
    return table


def _add_heading_styled(doc, text: str, level: int = 1):
    """添加带样式的标题（深蓝色，左侧竖线效果用缩进模拟）。"""
    from docx.shared import Pt, RGBColor
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = _TITLE_FONT
        run.font.color.rgb = RGBColor(0x0F, 0x3D, 0x6E)
        if level == 0:
            run.font.size = Pt(18)
        elif level == 1:
            run.font.size = Pt(13)
        elif level == 2:
            run.font.size = Pt(11.5)
    return h


def _add_body_para(doc, text: str):
    """添加正文段落（宋体 10.5pt, 首行缩进）。"""
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
    p = doc.add_paragraph(text)
    pf = p.paragraph_format
    pf.first_line_indent = Pt(21)  # 约两字符
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in p.runs:
        run.font.name = _BODY_FONT
        run.font.size = Pt(10.5)
    return p


def render_docx(context: dict) -> bytes:
    """Round10: 专业红头文件格式 DOCX — 封面页 + 页眉页脚 + 专业表格 + 水印。"""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    import base64 as _b64

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width  = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.top_margin    = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.left_margin   = Inches(0.98)
    section.right_margin  = Inches(0.98)

    # ── 页眉 ──
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("污染场地土壤生态-生产功能重构监管系统")
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    hr.font.name = _BODY_FONT

    # ── 页脚（页码） ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 插入页码域
    fr = fp.add_run()
    fldChar1 = fr._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    fr._element.append(fldChar1)
    instrText = fr._element.makeelement(qn("w:instrText"), {})
    instrText.text = " PAGE "
    fr._element.append(instrText)
    fldChar2 = fr._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    fr._element.append(fldChar2)
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── 水印（所有节） ──
    # 通过 sectPr 背景添加水印文字（VML）
    try:
        sectPr = section._sectPr
        vml = sectPr.makeelement(qn("w:background"), {})
        vml_v = vml.makeelement(qn("v:background"), {
            qn("v:fill"): "on",
            qn("v:fillcolor"): "#E0E4E8",
            qn("v:fillopacity"): ".15",
        })
        vml.append(vml_v)
        sectPr.insert(0, vml)
    except Exception:
        pass  # 水印非关键, 静默跳过

    # ═══════════════════════════════════════════════════
    # 封面页
    # ═══════════════════════════════════════════════════
    # 红色双线（模拟红头文件）
    for _ in range(8):
        doc.add_paragraph("")  # 空行推到中部

    # 红色双线
    red_line_p = doc.add_paragraph()
    red_line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl_run = red_line_p.add_run("━" * 44)
    rl_run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    rl_run.font.size = Pt(10)

    # 监管部门名称
    dept_p = doc.add_paragraph()
    dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dept_run = dept_p.add_run("生态环境部土壤与农业农村生态环境监管技术中心")
    dept_run.font.name = _TITLE_FONT
    dept_run.font.size = Pt(12)
    dept_run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)

    # 红色双线
    red_line_p2 = doc.add_paragraph()
    red_line_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl_run2 = red_line_p2.add_run("━" * 44)
    rl_run2.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    rl_run2.font.size = Pt(10)

    doc.add_paragraph("")

    # 报告标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("污染场地全流程监管追溯报告")
    title_run.font.name = _TITLE_FONT
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x0F, 0x3D, 0x6E)
    title_run.bold = True

    doc.add_paragraph("")

    # 封面信息表
    cover_info = [
        ("场地名称", context["site"]["name"]),
        ("场地编号", context["site"]["site_code"]),
        ("污染类型", context["site"]["pollution_type"] or "—"),
        ("用地类型", context["site"]["land_use_type"] or "—"),
        ("行政区划", f"{context['site']['province'] or ''} {context['site']['city'] or ''}"),
        ("报告版本", context["report"]["version"]),
        ("生成时间", context["report"]["generated_at"]),
        ("密级", "内部"),
    ]
    cover_table = doc.add_table(rows=len(cover_info), cols=2)
    cover_table.autofit = True
    for i, (k, v) in enumerate(cover_info):
        cover_table.rows[i].cells[0].text = k
        cover_table.rows[i].cells[1].text = v
    _style_table(cover_table, header_rows=0)
    for row in cover_table.rows:
        _set_cell_shading(row.cells[0], "F0F4F8")
        for p in row.cells[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 分页
    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # 正文内容
    # ═══════════════════════════════════════════════════
    _add_heading_styled(doc, "一、场地基本信息")
    _make_kv_table(doc, [
        ("场地编号", context["site"]["site_code"]),
        ("场地名称", context["site"]["name"]),
        ("污染类型", context["site"]["pollution_type"] or "—"),
        ("用地类型", context["site"]["land_use_type"] or "—"),
        ("行政区划", f"{context['site']['province'] or ''} {context['site']['city'] or ''}"),
        ("中心坐标", f"{context['site']['longitude'] or '—'}, {context['site']['latitude'] or '—'}"),
    ])

    _add_heading_styled(doc, "二、数据来源与覆盖率")
    _make_kv_table(doc, [
        ("来源文件", context["data_source"]["source_file"] or "—"),
        ("导入批次", f"#{context['data_source']['batch_id'] or '—'}（{context['data_source']['status'] or '—'}）"),
        ("采样点数", context["data_source"]["n_points"]),
        ("检测记录数", context["data_source"]["n_measurements"]),
        ("实测因子数", context["coverage"]["factor_count"]),
        ("覆盖率", f"{context['coverage']['coverage_pct']}%"),
        ("缺失率", f"{context['coverage']['missing_pct']}%"),
    ])

    # 地图图件
    _add_heading_styled(doc, "三、采样点空间分布")
    map_img = context["map_summary"].get("map_image")
    if map_img and map_img.startswith("data:image/png;base64,"):
        try:
            img_bytes = _b64.b64decode(map_img.split(",", 1)[1])
            from io import BytesIO as _BIO
            doc.add_picture(_BIO(img_bytes), width=docx_emu_width(doc))
            cp = doc.add_paragraph("▲ 采样点空间分布与超标风险分级（8级色阶, 离线渲染）")
            cp.italic = True
        except Exception:
            doc.add_paragraph("[地图图件渲染失败]")
    _make_kv_table(doc, [
        ("坐标覆盖", f"{context['map_summary']['n_coord_points']}/{context['map_summary']['n_points']} 个点位"),
        ("空间范围", str(context['map_summary']['bounds'] or "无")),
    ])

    # 检测数据摘要
    _add_heading_styled(doc, "四、检测数据摘要")
    if context.get("factor_summary"):
        table = doc.add_table(rows=1, cols=6)
        for i, h in enumerate(["因子", "类别", "样本数", "最小值", "均值", "最大值"]):
            table.rows[0].cells[i].text = h
        for f in context["factor_summary"][:20]:
            cells = table.add_row().cells
            cells[0].text = str(f["factor"])
            cells[1].text = str(f.get("category") or "")
            cells[2].text = str(f["count"])
            cells[3].text = str(f["min"])
            cells[4].text = str(f["mean"])
            cells[5].text = str(f["max"])
        _style_table(table)
    else:
        doc.add_paragraph("暂无检测数据。")
    _embed_docx_image(doc, _render_eda_figure_png(context.get("factor_summary") or []),
                      "▲ 各因子浓度均值与最大值对比")

    # 数据质量校验
    _add_heading_styled(doc, "五、数据质量校验")
    _make_kv_table(doc, [
        ("校验结论", "✓ 通过" if context["validation"]["passed"] else "✗ 存在阻断性错误"),
        ("错误/警告", f"{context['validation']['n_errors']}/{context['validation']['n_warnings']}"),
        ("超标指标", f"{context['validation']['n_exceed']} 项: {'、'.join(context['validation']['exceed_factors']) or '无'}"),
    ])

    # 障碍因子诊断
    _add_heading_styled(doc, "六、障碍因子诊断")
    if context.get("diagnosis") and context["diagnosis"].get("top_factors"):
        table = doc.add_table(rows=1, cols=5)
        for i, h in enumerate(["排名", "因子", "类别", "KOS/贡献值", "方向"]):
            table.rows[0].cells[i].text = h
        for t in context["diagnosis"]["top_factors"]:
            cells = table.add_row().cells
            cells[0].text = str(t["rank"])
            cells[1].text = str(t["factor"])
            cells[2].text = str(t.get("category") or "")
            cells[3].text = str(t.get("importance", t.get("kos_score", t.get("KOS", ""))))
            cells[4].text = str(t.get("direction") or "")
        _style_table(table)
    else:
        doc.add_paragraph("暂无诊断结果。")
    _embed_docx_image(doc, context["map_summary"].get("shap_image"),
                      "▲ 关键障碍因子模型贡献份额排名")

    # 功能重构
    _add_heading_styled(doc, "七、功能重构可行性评价")
    if context.get("reconstruction"):
        for ev in context["reconstruction"]:
            _add_body_para(doc,
                f"{ev['title']}: 得分 {ev['score']} ({ev['grade']}), "
                f"限制因子: {'、'.join(ev.get('limiting_factors') or []) or '无'}")
            if ev.get("explanation"):
                _add_body_para(doc, str(ev["explanation"]))
    else:
        doc.add_paragraph("暂无功能重构评价结果。")

    # SSUI
    _add_heading_styled(doc, "八、可持续利用评价（SSUI）")
    if context.get("ssui"):
        _make_kv_table(doc, [
            ("SSUI 指数", context["ssui"]["score"]),
            ("可持续性等级", context["ssui"]["grade"]),
            ("说明", str(context["ssui"].get("explanation") or "—")),
        ])
    else:
        doc.add_paragraph("暂无 SSUI 结果。")

    # 推荐方案
    _add_heading_styled(doc, "九、推荐修复方案矩阵")
    if context.get("recommendations"):
        table = doc.add_table(rows=1, cols=6)
        for i, h in enumerate(["排序", "技术", "匹配度", "成本", "禁用条件", "理由"]):
            table.rows[0].cells[i].text = h
        for r in context["recommendations"]:
            cells = table.add_row().cells
            cells[0].text = str(r["rank"])
            cells[1].text = str(r["technology"])
            cells[2].text = str(r["match_score"])
            cells[3].text = str(r.get("cost_level") or "")
            cells[4].text = str(r.get("forbidden_conditions") or "")
            cells[5].text = (str(r.get("reason") or ""))[:240]
        _style_table(table)
    else:
        doc.add_paragraph("暂无推荐方案。")

    # 修复案例
    _add_heading_styled(doc, "十、修复案例证据库")
    for c in context.get("remediation_cases", [])[:6]:
        doc.add_paragraph(
            f"{c['case_id']}｜{c['remediation_technology']}｜{c['pollutants']}｜"
            f"证据: {c.get('evidence_source') or '—'}"
        )

    # 追溯记录
    _add_heading_styled(doc, "十一、五阶段全流程追溯记录")
    if context.get("workflow"):
        table = doc.add_table(rows=1, cols=5)
        for i, h in enumerate(["阶段", "状态", "版本", "审批意见", "附件数"]):
            table.rows[0].cells[i].text = h
        for w in context["workflow"]:
            cells = table.add_row().cells
            cells[0].text = str(w["stage_name"])
            cells[1].text = str(w["status"])
            cells[2].text = str(w.get("version") or "")
            cells[3].text = str(w.get("review_comment") or "")
            cells[4].text = str(w.get("n_attachments") or 0)
        _style_table(table)
    else:
        doc.add_paragraph("暂无追溯记录。")

    # 附件
    _add_heading_styled(doc, "十二、附件清单")
    if context.get("attachments"):
        for a in context["attachments"]:
            doc.add_paragraph(f"{a['stage_name']}｜{a.get('file_role') or '—'}｜{a['original_name']}")
    else:
        doc.add_paragraph("暂无附件。")

    # 版本信息
    _add_heading_styled(doc, "十三、版本与口径说明")
    _make_kv_table(doc, [
        ("模型版本", context["diagnosis"]["model_version"] if context.get("diagnosis") else "—"),
        ("数据版本", context["report"]["data_version"]),
        ("标准版本", context["report"]["standard_version"]),
        ("模板版本", context["report"]["template_version"]),
        ("报告版本", context["report"]["version"]),
    ])

    _add_heading_styled(doc, "报告口径与结果范围说明")
    _add_body_para(doc,
        "正式超标结论仅限于身份明确、单位兼容且阈值适用的因子。"
        "对没有适用阈值或未被正式因子库收录的实测指标，系统仍通过模型候选识别、"
        "族群级近邻分析和未知因子预警进行辅助识别，不会丢弃数据或强行套用标准。"
        "探索性识别结果不等同于法规超标判定，需结合检测方法和专家复核。")
    _add_body_para(doc,
        "当前完成 3 个原始场地的工程回归验证 + 15 个合成数据演示，"
        "尚未开展跨区域大规模独立验证。报告结论作为辅助决策依据，"
        "不构成监管级科学可信判定。")
    _add_body_para(doc,
        "采用规则、模型和开放集识别相结合的混合策略(非纯数据驱动)。"
        "AI 润色文本经事实校验但仍有降级回退机制；"
        "任何 AI 生成的描述均以原始检测数据为准。")

    # 人工复核区
    doc.add_paragraph("")
    _add_heading_styled(doc, "人工复核意见")
    doc.add_paragraph("（请在此处填写复核意见）")
    doc.add_paragraph("\n\n")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Round10: 报告类型中文标签
REPORT_SCOPE_LABEL = {
    "full": "全流程追溯报告",
    "ssui": "SSUI可持续利用评价报告",
    "diagnosis": "障碍因子诊断报告",
    "reconstruction": "功能重构评价报告",
}


def generate(db: Session, site_id: int, generated_by: int | None = None,
             report_format: str = "pdf", report_scope: str = "full") -> dict:
    n_prev = db.query(ReportRecord).filter_by(site_id=site_id).count()
    version = f"v{n_prev + 1}"
    ctx = collect(db, site_id, version)
    requested = (report_format or "pdf").lower()
    scope = (report_scope or "full").lower()
    scope_label = REPORT_SCOPE_LABEL.get(scope, "全流程追溯报告")
    report_type = scope  # 存入 ReportRecord.report_type

    if requested == "docx":
        docx_bytes = render_docx(ctx)
        fo = save_bytes(
            db, docx_bytes, f"{scope_label}_{ctx['site']['site_code']}_{version}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        fmt = "docx"
    elif requested == "html":
        html = render_html(ctx)
        fo = save_bytes(db, html.encode("utf-8"),
                        f"{scope_label}_{ctx['site']['site_code']}_{version}.html",
                        content_type="text/html")
        fmt = "html"
    else:
        html = render_html(ctx)
        pdf = html_to_pdf(html)
        if pdf:
            fo = save_bytes(db, pdf, f"{scope_label}_{ctx['site']['site_code']}_{version}.pdf",
                            content_type="application/pdf")
            fmt = "pdf"
        else:
            fo = save_bytes(db, html.encode("utf-8"),
                            f"{scope_label}_{ctx['site']['site_code']}_{version}.html",
                            content_type="text/html")
            fmt = "html"

    rec = ReportRecord(
        site_id=site_id, report_type=report_type, version=version,
        data_snapshot={"data_version": ctx["report"]["data_version"],
                       "standard_version": ctx["report"]["standard_version"],
                       "format": fmt,
                       "scope": scope,
                       "diagnosis": bool(ctx["diagnosis"]),
                       "n_recommendations": len(ctx["recommendations"]),
                       "n_remediation_cases": len(ctx["remediation_cases"]),
                       "validation_passed": ctx["validation"]["passed"]},
        template_version=TEMPLATE_VERSION, file_object_id=fo.id,
        generated_by=generated_by, generated_at=datetime.now(timezone.utc))
    db.add(rec)
    db.commit()
    return {"report_id": rec.id, "site_id": site_id, "version": version,
            "format": fmt, "scope": scope, "file_object_id": fo.id,
            "storage_key": fo.storage_key, "file_name": fo.original_name}
