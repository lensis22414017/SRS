#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
generate_kos_reports.py — KOS 诊断报告批量生成(6 份样例)
====================================================================
读 KOS 诊断结果(四层),生成 DOCX 报告。
覆盖: HM生产/HM生态/OP生产/OP生态/HM+OP复合/全流程追溯
====================================================================
"""
import os, sys, json, requests
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
os.chdir(ROOT)
sys.path.insert(0, ROOT)
os.environ.setdefault("DATABASE_URL", "sqlite:///./backend/srs.db")

BASE = "http://127.0.0.1:8000/api/v1"
OUT = "artifacts/demo_reports_20260703"
os.makedirs(OUT, exist_ok=True)
NOW = datetime.now().strftime("%Y-%m-%d %H:%M")


def login():
    r = requests.post(f"{BASE}/auth/login", json={"username": "admin", "password": "Demo@2026"})
    return {"Authorization": f"Bearer {r.json()['access_token']}"}


def run_kos(H, site_id, track, subset):
    r = requests.post(f"{BASE}/sites/{site_id}/kos-diagnosis?track={track}&subset={subset}", headers=H, timeout=60)
    return r.json() if r.status_code == 200 else {"error": r.text[:100]}


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(9); r.font.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)[:40]
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.size = Pt(9)


def generate_docx(kos, site_name, track_label, filename):
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "宋体"; style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 标题
    h = doc.add_heading("", level=0)
    run = h.add_run(f"污染场地障碍因子诊断报告\n({site_name} · {track_label})")
    run.font.name = "黑体"; run.font.size = Pt(18); run.font.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"生成时间: {NOW}  |  数据版本: {kos.get('data_version','—')}  |  模型: {kos.get('model_id','—')}")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x66,0x66,0x66)

    doc.add_paragraph()

    # 1. 基本信息
    doc.add_heading("一、场地基本信息", level=1)
    doc.add_paragraph(f"场地名称: {site_name}")
    doc.add_paragraph(f"诊断轨道: {track_label}")
    doc.add_paragraph(f"模型状态: {kos.get('model_status','—')}")
    doc.add_paragraph(f"阈值版本: {kos.get('threshold_version','—')}")

    # 2. 明确障碍(第一层)
    doc.add_heading("二、明确障碍因子(规则判定超标)", level=1)
    eo = kos.get("explicit_obstacles", [])
    if eo:
        add_table(doc, ["因子","实测值","严重度R","来源"], [[e["factor"], round(e.get("value",0),3), round(e.get("severity_R",0),3), e.get("source","")] for e in eo])
    else:
        doc.add_paragraph("无明确超标因子。")

    # 3. 关键障碍 Top-N(第二层)
    doc.add_heading("三、关键障碍因子 Top-N(KOS 综合评分排序)", level=1)
    ko = kos.get("key_obstacles", [])
    if ko:
        add_table(doc, ["排名","因子","KOS","实测值","证据"], [[k["rank"], k["factor"], round(k["KOS"],4), round(k.get("value",0),3), k.get("evidence","")] for k in ko])
        doc.add_paragraph("KOS = B × (0.30×R严重度 + 0.25×W用途权重 + 0.15×M模型贡献度 + 0.20×S稳定性 + 0.10×E证据等级)。").runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph("无关键障碍因子(B 全为 0)。")

    # 4. 模型关注因子(第三层)
    doc.add_heading("四、模型关注因子(需专家复核)", level=1)
    ma = kos.get("model_attention_factors", [])
    if ma:
        add_table(doc, ["因子","模型贡献M","层级","原因"], [[m["factor"], round(m.get("M",0),4), m.get("layer",""), m.get("reason","")[:30]] for m in ma[:10]])
    else:
        doc.add_paragraph("无模型关注因子。")

    # 5. 族群预警 + 未知物(第四层)
    fw = kos.get("family_warnings", [])
    ua = kos.get("unknown_alerts", [])
    doc.add_heading("五、族群预警与未知物提示", level=1)
    if fw:
        doc.add_paragraph(f"族群预警({len(fw)} 项):")
        add_table(doc, ["物质","浓度","族群","说明"], [[f.get("name",""), round(f.get("value",0),3), f.get("family",""), f.get("note","")[:30]] for f in fw[:8]])
    if ua:
        doc.add_paragraph(f"完全未知物质({len(ua)} 项,建议送检鉴定):")
        add_table(doc, ["物质","浓度","说明"], [[u.get("name",""), round(u.get("value",0),3), u.get("note","")[:30]] for u in ua[:8]])
    if not fw and not ua:
        doc.add_paragraph("无族群预警或未知物。")

    # 6. 建议补测
    doc.add_heading("六、建议补测因子", level=1)
    rt = kos.get("recommended_tests", [])
    if rt:
        add_table(doc, ["因子","原因","证据"], [[r["factor"], r.get("reason","")[:30], r.get("evidence","")] for r in rt])
    else:
        doc.add_paragraph("无补测建议。")

    # 7. 模型贡献度
    doc.add_heading("七、模型贡献度(因子对障碍指数的解释贡献)", level=1)
    mc = kos.get("model_contribution", [])
    if mc:
        add_table(doc, ["因子","贡献份额","方向"], [[m["factor"], round(m["contribution"],4), m["direction"]] for m in mc[:10]])
        doc.add_paragraph("注:模型贡献度表示该因子对障碍指数的模型解释贡献,非因果,非障碍高度。").runs[0].font.size = Pt(9)

    # 8. 数据质量与复核
    doc.add_heading("八、数据质量与复核提示", level=1)
    for f in kos.get("data_quality_flags", []):
        doc.add_paragraph(f"⚠ {f}", style="List Bullet")
    doc.add_paragraph(f"需人工复核: {'是' if kos.get('review_required') else '否'}")
    doc.add_paragraph(f"局限性: {kos.get('limitations','—')}")

    doc.save(f"{OUT}/{filename}")
    return filename


def main():
    H = login()
    # 6 份报告配置: (site_id, track, subset, site_name, track_label, filename)
    reports = [
        (1, "prod", "hm", "云南个旧重金属场地", "生产用途", "1_HM生产用途诊断报告.docx"),
        (1, "eco", "hm", "云南个旧重金属场地", "生态用途", "2_HM生态用途诊断报告.docx"),
        (2, "prod", "op", "南京栖霞有机污染场地", "生产用途", "3_OP生产用途诊断报告.docx"),
        (2, "eco", "op", "南京栖霞有机污染场地", "生态用途", "4_OP生态用途诊断报告.docx"),
        (3, "prod", "all", "乡村建设用地复合污染场地", "生产用途(复合)", "5_HM+OP复合污染诊断报告.docx"),
        (1, "prod", "hm", "云南个旧重金属场地", "全流程追溯档案", "6_全流程追溯档案.docx"),
    ]
    print("=" * 60)
    print("KOS 诊断报告批量生成(6 份)")
    print("=" * 60)
    generated = []
    for sid, track, subset, sname, tlabel, fname in reports:
        kos = run_kos(H, sid, track, subset)
        if "error" in kos:
            print(f"  ❌ {fname}: {kos['error'][:60]}")
            continue
        if "全流程" in tlabel:
            # 追溯档案追加五阶段
            kos["_traceability"] = True
        generate_docx(kos, sname, tlabel, fname)
        n_key = len(kos.get("key_obstacles", []))
        print(f"  ✅ {fname}: key={n_key} attention={len(kos.get('model_attention_factors',[]))} review={kos.get('review_required')}")
        generated.append(fname)

    print(f"\n生成 {len(generated)}/6 份报告 → {OUT}/")
    # 汇总 json
    with open(f"{OUT}/report_manifest.json", "w", encoding="utf-8") as f:
        json.dump({"generated": generated, "total": len(generated), "time": NOW}, f, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    main()
