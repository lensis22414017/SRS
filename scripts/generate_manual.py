"""生成 SRS 系统用户操作手册 (DOCX 标准公文格式)"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)


def add_title(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
    return h


def add_body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_item(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p


# ==== 封面 ====
for _ in range(4):
    doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('污染场地土壤生态-生产功能'); r.font.name = '黑体'; r.font.size = Pt(26); r.bold = True
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('障碍识别与重构利用监管系统'); r.font.name = '黑体'; r.font.size = Pt(26); r.bold = True
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('用户操作手册'); r.font.name = '黑体'; r.font.size = Pt(18)
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('版本 V1.0.1'); r.font.size = Pt(14)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('2026年7月'); r.font.size = Pt(14)
doc.add_page_break()

# ==== 第一章 ====
add_title('第一章  系统概述', 1)
add_title('1.1  系统简介', 2)
add_body('污染场地土壤生态-生产功能障碍识别与重构利用监管系统（以下简称"本系统"）是一套面向污染场地修复治理全流程的监管平台。系统基于《污染场地土壤生态-生产功能障碍识别与重构利用的评价方法》开发，实现了从数据导入、障碍因子诊断、功能重构可行性评价、可持续利用评价到修复方案推荐和全流程追溯的完整业务闭环。')

add_title('1.2  主要功能模块', 2)
for item in [
    '(1) 数据概览大屏：展示场地总数、检测记录、评价进度等全局统计信息。',
    '(2) 场地管理：支持场地信息录入、采样点管理、检测数据导入与导出。',
    '(3) 障碍因子诊断：基于规则与模型混合策略，识别场地关键障碍因子。',
    '(4) 功能重构评价：评估污染场地修复后的生产功能和生态功能重构可行性。',
    '(5) 可持续利用评价（SSUI）：综合25项指标评估场地可持续利用水平。',
    '(6) 修复方案推荐：基于场地特征匹配推荐适用的修复技术和方案。',
    '(7) 全流程追溯：五阶段（调查—审批—施工—效果—管护）进度管理与文件归档。',
    '(8) 报告生成：支持 PDF、DOCX、HTML 三种格式的全流程追溯报告导出。',
]:
    add_item(item)

# ==== 第二章 ====
add_title('第二章  安装与启动', 1)
add_title('2.1  系统要求', 2)
for req in ['操作系统：Windows 10/11（64位）。', '内存：建议 8GB 及以上。', '硬盘空间：至少 2GB 可用空间。', '屏幕分辨率：建议 1920×1080 及以上。']:
    add_item(req)

add_title('2.2  安装步骤', 2)
for step in ['(1) 运行安装程序 SRS-Setup-V1.0.1-Windows-x64.exe。', '(2) 按照安装向导提示选择安装路径。', '(3) 安装完成后，桌面将生成系统快捷方式。', '(4) 双击快捷方式启动系统，浏览器自动打开访问界面。']:
    add_item(step)

add_title('2.3  首次使用', 2)
for step in ['(1) 首次启动后系统自动打开浏览器访问 http://localhost:8000。', '(2) 首次登录需设置管理员账号和密码。', '(3) 系统预置30个演示场地数据，可直接用于功能体验和测试。', '(4) 演示数据中的经济指标采用全国统计数据作为参考值，实际使用时请替换为场地真实数据。']:
    add_item(step)

# ==== 第三章 ====
add_title('第三章  操作指南', 1)

add_title('3.1  数据导入', 2)
add_body('系统支持 Excel（.xlsx/.xls）和 CSV 格式的检测数据导入。操作步骤如下：')
for s in [
    '(1) 点击左侧导航栏「数据导入」进入导入页面。',
    '(2) 选择字段映射模板（推荐使用「自动识别」）。',
    '(3) 选择重复导入策略（跳过/覆盖/新版本）。',
    '(4) 拖拽或点击上传文件，支持批量多文件同时导入。',
    '(5) 点击「开始导入并校验」，系统自动解析、映射、校验并写入数据库。',
    '(6) 导入完成后可查看导入质量报告，含校验错误和超标因子统计。',
]:
    add_item(s)

add_title('3.2  障碍因子诊断', 2)
add_body('障碍因子诊断模块用于识别场地中超过国家土壤污染风险管控标准的关键障碍因子。操作步骤如下：')
for s in [
    '(1) 在顶部选择目标场地。',
    '(2) 选择修复后用途（生产用地/生态用地），系统自动切换对应评价轨。',
    '(3) 点击「运行障碍因子诊断」按钮。',
    '(4) 系统输出关键障碍因子 Top-N 排名、KOS 综合评分及五分量证据堆叠图。',
    '(5) 可查看历史诊断记录，支持多版本对比分析。',
]:
    add_item(s)

add_title('3.3  功能重构评价', 2)
add_body('功能重构评价分为生产功能重构和生态功能重构两个维度，采用改进模糊综合评价法。操作步骤如下：')
for s in [
    '(1) 选择场地后点击「运行功能重构可行性评价」。',
    '(2) 系统分别输出生产功能和生态功能重构的综合得分与等级。',
    '(3) 雷达图展示各评价指标得分，条形图展示指标贡献度排序。',
    '(4) 瀑布图展示累计贡献叠加，仪表盘展示最短板指标。',
    '(5) 有机污染场地自动展示有机污染风险诊断作为替代分析。',
]:
    add_item(s)

add_title('3.4  SSUI 可持续利用评价', 2)
add_body('SSUI（Soil Sustainable Utilization Index，土壤可持续利用综合指数）基于25项元指标（D1-D25）加权计算。操作步骤如下：')
for s in [
    '(1) 选择场地，设置评价参数（利用年限1-50年、管理强度粗放/中等/集约、评价年份）。',
    '(2) 如需使用全国平均经济数据代替场地真实数据，勾选对应复选框。',
    '(3) 点击「运行评价」按钮。',
    '(4) 系统输出 SSUI 指数、可持续性等级、D1-D25 元指标得分明细。',
    '(5) 仪表盘展示综合指数（0-1范围），二维象限图展示安全性与经济性平衡关系。',
    '(6) D1-D15为土壤安全性指标（采用外部参照归一化），D16-D17为污染物风险指标（采用国家法规阈值），D18-D25为经济指标（采用版本化官方年度参照数据）。',
]:
    add_item(s)

add_title('3.5  经济数据管理', 2)
add_body('SSUI 评价依赖 D18-D25 共8项经济指标。系统提供经济数据管理功能：')
for s in [
    '(1) 在 SSUI 页面点击「经济数据」按钮打开管理面板。',
    '(2) 支持逐项录入或通过 Excel 模板批量导入。',
    '(3) 可下载标准模板（含指标定义、单位、方向说明）。',
    '(4) 系统内置全国统计数据作为参考值，来源包括国家统计局年度公报、全国农产品成本收益资料汇编等。',
    '(5) 使用参考数据时评价结果标记为「参考评价」，正式评价需录入场地真实经营数据。',
]:
    add_item(s)

add_title('3.6  全流程追溯', 2)
add_body('全流程追溯模块覆盖污染场地修复治理的五阶段管理：')
for s in [
    '调查阶段：上传场地调查报告、检测数据、障碍因子识别结果等。',
    '审批阶段：上传重构方案、修改记录、最终通过版本等。',
    '施工阶段：上传施工方案、监理方案、施工进度记录、材料使用台账等。',
    '效果评估阶段：上传效果检测数据、效果评估报告、达标结论等。',
    '管护阶段：上传管护方案、定期监测数据、功能维护记录等。',
]:
    add_item(s)
add_body('每个阶段支持文件上传、下载、预览和删除操作。系统自动计算证据链完整度，标识材料缺口。点击「生成报告」可导出全流程 PDF/DOCX/HTML 格式追溯报告。')

# ==== 第四章 ====
add_title('第四章  报告生成', 1)
add_body('系统支持从任意分析页面生成综合报告，包含场地基本信息、检测数据摘要、障碍因子排名、功能重构评价、SSUI 评价、修复方案推荐、五阶段追溯记录等内容。点击各分析页面右上角「导出报告」按钮即可生成并下载。报告支持三种输出格式：PDF（默认，适合打印存档）、DOCX（适合后续编辑）、HTML（适合在线浏览）。')

# ==== 第五章 ====
add_title('第五章  评价体系说明', 1)
add_title('5.1  D1-D25 元指标体系', 2)
add_body('本系统采用25项元指标（D1-D25）构成四层评价体系。目标层A为土壤可持续利用综合指数（SSUI）。准则层B1（安全性）包含限制因子C1（D1-D15，15项土壤基础指标）和风险因子C2（D16-D17，重金属与有机污染物）。准则层B2（经济性）包含经济成本C3（D18-D21，劳动力、机械、土地、非机械化生产成本）和经济效益C4（D22-D25，总产值、效益费用比、人均可支配收入、实物产量）。')
add_body('计算公式：SSUI = (B1 × 0.5 + B2 × 0.5) × f(t) × M，其中 f(t) = 1 + 0.03 × t 为时间修正函数（t为修复后利用年数），M 为管理调节因子（根据粗放/中等/集约管理强度分别取值）。')

add_title('5.2  可持续性等级划分', 2)
for s in [
    '高可持续性：SSUI ≥ 0.80',
    '中高可持续性：0.60 ≤ SSUI < 0.80',
    '中可持续性：0.40 ≤ SSUI < 0.60',
    '低可持续性：SSUI < 0.40',
]:
    add_item(s)

add_title('5.3  正式评价与参考评价', 2)
add_body('正式评价要求 D1-D25 全部25项指标具备可审计的实测数据。当部分指标缺失或采用全国参照数据时，系统生成参考评价并明确标注。参考评价结论仅供初步参考，不作为场地正式结论。对于纯有机污染场地，由于当前评价体系基于重金属和农业肥力指标设计，系统提供有机污染风险诊断作为替代分析。')

# ==== 第六章 ====
add_title('第六章  常见问题与故障排除', 1)
faqs = [
    ('问：为什么某些场地显示「评价受阻」？',
     '答：当25项指标中关键指标缺失时（如缺少经济数据、土壤基础指标等），系统无法生成正式评价。请检查并补齐缺失数据后重试。'),
    ('问：经济数据如何获取？',
     '答：系统内置了国家统计局公布的全国平均参照数据（2015-2024年）。可通过「经济数据」管理面板录入场地真实经营数据，或导入 Excel 模板。'),
    ('问：有机污染场地为什么显示为不适用？',
     '答：当前 SSUI 评价体系基于重金属和农业肥力指标设计，纯有机污染场地缺少体系内评价所需的指标数据。系统会提供有机污染风险诊断作为替代分析。'),
    ('问：报告生成失败怎么办？',
     '答：请确认系统已安装完整依赖。如 PDF 生成失败，系统会自动降级为 HTML 格式输出，内容完整不丢失。'),
    ('问：如何添加新用户？',
     '答：在系统管理页面可进行用户注册与审批管理。新用户注册后需管理员审批通过方可登录使用。'),
]
for q, a in faqs:
    add_item(q)
    add_item(a)
    doc.add_paragraph()

# ==== 附录 ====
add_title('附录  技术规格', 1)
add_title('A.1  数据标准', 2)
for s in [
    '检测数据标准：GB 15618-2018《土壤环境质量 农用地土壤污染风险管控标准》。',
    '建设用地标准：GB 36600-2018《土壤环境质量 建设用地土壤污染风险管控标准》。',
    '评价方法依据：《污染场地土壤生态-生产功能障碍识别与重构利用的评价方法》（2025年版）。',
    '经济数据来源：国家统计局年度公报、全国农产品成本收益资料汇编。',
]:
    add_item(s)

add_title('A.2  技术架构', 2)
for s in [
    '后端技术：Python 3.11+ / FastAPI / SQLite。',
    '前端技术：React 18 / TypeScript / Ant Design / ECharts。',
    '机器学习框架：Scikit-learn / SHAP / XGBoost。',
    '报告引擎：Jinja2 / WeasyPrint / python-docx。',
]:
    add_item(s)

out_dir = r'C:\Users\曾鸿\Desktop\SRS-round10\交付物'
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, 'SRS用户操作手册_V1.0.1.docx')
doc.save(out_path)
print(f'DOCX generated: {out_path} ({os.path.getsize(out_path)/1024:.0f} KB)')
