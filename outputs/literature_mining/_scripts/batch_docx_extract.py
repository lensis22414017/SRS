"""批量提取 HM+OP 候选 SI docx 表格 → markdown, 扫描含采样点浓度表的 SI

裴总铁律: 脚本只定位(哪些SI有采样点表格), 不提取(提取靠Agent精读)
扫描: 表格行≥8 + 含浓度单位(mg/kg|ng/g) + 行/列含采样点标识 + 非参数表(RfD/SF/IngR暴露参数)
"""
from __future__ import annotations
import sys, re
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

import pandas as pd

ROOT = Path(r"G:\文献整理_最终")
OUT = Path(r"C:\Users\曾鸿\Desktop\SRS\outputs\literature_mining")
SI_OUT = OUT / "si_extract"
SI_OUT.mkdir(exist_ok=True)

UNIT_RE = re.compile(r"mg\s*/?\s*kg|μg|ug\s*/?\s*kg|ng\s*/?\s*g|ppm|ppb|mg·kg", re.I)
# 暴露参数/风险参数表关键词 (排除)
PARAM_RE = re.compile(r"RfD|\bSF\b|IngR|InhR|\bEF\b|\bED\b|\bAT\b|\bBW\b|\bIR\b|CSF|PEF|\bABS\b|\bAF\b|exposure\s*factor|risk\s*param|Tef|TEF|toxic\s*equival", re.I)
SAMPLE_LABEL_RE = re.compile(r"^S\d{1,4}$|^sample\s*\d|^site\s*\d|^station|^[A-Z]\d{0,2}$|^\d{1,4}$|点位|采样", re.I)


def table_to_md(table) -> str:
    lines = []
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def classify_table(table) -> dict:
    rows = table.rows
    n_rows = len(rows)
    if n_rows < 4:
        return {"n_rows": n_rows, "has_unit": False, "is_param": False, "sample_signals": 0}
    text = " ".join(c.text for r in rows for c in r.cells)
    has_unit = bool(UNIT_RE.search(text))
    is_param = bool(PARAM_RE.search(text)) and not bool(re.search(r"\bS\d+\b|sample\s*\d", text, re.I))
    # 采样点信号: 第一列(行标签)或表头列名像采样点
    sample_cnt = 0
    for r in rows[1:]:
        if r.cells:
            lbl = r.cells[0].text.strip()
            if SAMPLE_LABEL_RE.match(lbl):
                sample_cnt += 1
    # 表头列名采样点
    header = [c.text.strip() for c in rows[0].cells] if rows[0].cells else []
    col_samples = sum(1 for c in header[1:] if SAMPLE_LABEL_RE.match(c))
    return {"n_rows": n_rows, "has_unit": has_unit, "is_param": is_param,
            "sample_signals": max(sample_cnt, col_samples)}


def process_docx(pid: str, stem: str):
    si_dir = ROOT / stem / "si"
    if not si_dir.exists():
        return None
    docx_files = [f for f in si_dir.glob("*.docx") if not f.name.startswith("._")]
    if not docx_files:
        return None
    try:
        import docx
    except ImportError:
        return None
    all_md = []
    best = {"n_tables": 0, "max_rows": 0, "sample_tables": 0, "best_sample_signals": 0}
    for df in docx_files:
        try:
            doc = docx.Document(str(df))
        except Exception:
            continue
        best["n_tables"] += len(doc.tables)
        for i, tbl in enumerate(doc.tables):
            info = classify_table(tbl)
            best["max_rows"] = max(best["max_rows"], info["n_rows"])
            # 采样点表: 行≥8 + 有单位 + 非参数 + 采样点信号≥5
            if info["n_rows"] >= 8 and info["has_unit"] and not info["is_param"] and info["sample_signals"] >= 5:
                best["sample_tables"] += 1
                best["best_sample_signals"] = max(best["best_sample_signals"], info["sample_signals"])
            all_md.append(f"\n## {df.name} - Table {i+1} ({info['n_rows']}行)\n")
            all_md.append(table_to_md(tbl))
    # 写 markdown
    if all_md:
        (SI_OUT / f"{pid}_si.md").write_text("\n".join(all_md), encoding="utf-8")
    best["paper_id"] = pid
    best["stem"] = stem
    return best


def main():
    scan = pd.read_csv(OUT / "scan_sample_row_hmop.csv", dtype=str, keep_default_na=False)
    print(f"批量提取 HM+OP SI docx ({len(scan)} 候选)...")
    results = []
    for i, r in scan.iterrows():
        info = process_docx(r["paper_id"], r["stem"])
        if info:
            results.append(info)
        if (i + 1) % 100 == 0:
            print(f"  进度 {i+1}/{len(scan)}")
    df = pd.DataFrame(results)
    df.to_csv(OUT / "si_docx_scan.csv", index=False, encoding="utf-8-sig")

    cand = df[df["sample_tables"] >= 1]
    print(f"\n=== SI docx 扫描结果 ===")
    print(f"有SI docx: {len(df)} 篇")
    print(f"SI含采样点表格候选 (行≥8+单位+非参数+采样信号≥5): {len(cand)} 篇")
    print(f"\n候选 SI (按采样信号降序):")
    for _, r in cand.sort_values("best_sample_signals", ascending=False).head(25).iterrows():
        print(f"  {r['paper_id']} | {int(r['sample_tables'])}采样表/{int(r['n_tables'])}总表 | max{int(r['max_rows'])}行 sig{int(r['best_sample_signals'])} | {r['stem'][:38]}")


if __name__ == "__main__":
    main()
