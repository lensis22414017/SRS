"""输出DOCX表格尺寸与前几行，供SI提取器设计使用。"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("paths", nargs="+")
    parser.add_argument("--rows", type=int, default=5)
    args = parser.parse_args()
    for value in args.paths:
        path = Path(value)
        document = Document(path)
        print(f"\nFILE {path} tables={len(document.tables)}")
        for index, table in enumerate(document.tables):
            width = max((len(row.cells) for row in table.rows), default=0)
            print(f"TABLE {index}: rows={len(table.rows)} cols={width}")
            for row in table.rows[: args.rows]:
                print(" | ".join(cell.text.replace("\n", " ").strip() for cell in row.cells)[:2000])


if __name__ == "__main__":
    main()
