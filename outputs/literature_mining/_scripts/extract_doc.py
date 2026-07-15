import sys, io, csv, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document

doc = Document(r'C:\temp\P07060_SI.docx')

paper_id = 'P07060'

# Compound order in tables
compounds = ['ACY', 'ACE', 'FLO', 'PHE', 'ANT', 'FLA', 'PYR', 'BaA', 'CHR', 'BbF', 'BkF', 'BaP', 'IcdP', 'DahA', 'BghiP', 'PAHs']

all_data = []

# Helper to clean value
def parse_value(v):
    v = v.strip()
    if not v or v in ['-', 'BDL', '']:
        return None
    # Handle "BDL 1" (from Table SI-8 BbF row)
    v = v.replace('BDL ', '').replace('BDL', '').strip()
    if not v:
        return None
    # Remove ±stddev
    v = v.split('±')[0].strip()
    try:
        return float(v)
    except ValueError:
        return None

# Table SI-6 (index 6): Soil zone II, 14 sites + average
tbl6 = doc.tables[6]
# Row 1: site names
si6_sites = []
for c in tbl6.rows[1].cells[1:]:
    n = c.text.strip()
    si6_sites.append(n)
# Filter 'average' - keep only individual sites
si6_individual = [s for s in si6_sites if s and s.lower() != 'average']
print(f'SI-6 individual sites: {si6_individual}')

# Row 2-17: ACY-PAHs
for ri in range(2, 18):
    comp = tbl6.rows[ri].cells[0].text.strip()
    if not comp:  # Missing compound name (FLA row)
        # Use known compound order
        row_idx = ri - 2  # 0-based compound index
        if row_idx < len(compounds):
            comp = compounds[row_idx]
        else:
            continue

    # Normalize compound name
    matched = None
    for c in compounds:
        if c in comp or c.lower() in comp.lower():
            matched = c
            break
    if not matched:
        continue

    for si, site in enumerate(si6_sites):
        if site.lower() == 'average':
            continue
        if si + 1 >= len(tbl6.rows[ri].cells):
            break
        val = parse_value(tbl6.rows[ri].cells[si + 1].text)
        if val is not None and matched == 'PAHs':
            all_data.append([paper_id, site, 'Sum_PAH_ngg', val, 'ng/g', 'Table SI-6', 'soil', 'other', 'Tibet',
                           f'15 EPA PAH monomer sum from Table SI-6'])

# Table SI-7 (index 7): Soil zones III + IV
tbl7 = doc.tables[7]
si7_sites = []
for c in tbl7.rows[1].cells[1:]:
    n = c.text.strip()
    si7_sites.append(n)
si7_individual = [s for s in si7_sites if s and s.lower() != 'average' and s]
print(f'SI-7 individual sites: {si7_individual}')

for ri in range(2, 18):
    comp = tbl7.rows[ri].cells[0].text.strip()
    if not comp:
        row_idx = ri - 2
        if row_idx < len(compounds):
            comp = compounds[row_idx]
        else:
            continue
    matched = None
    for c in compounds:
        if c in comp or c.lower() in comp.lower():
            matched = c
            break
    if not matched:
        continue
    for si, site in enumerate(si7_sites):
        if site.lower() == 'average':
            continue
        if si + 1 >= len(tbl7.rows[ri].cells):
            break
        val = parse_value(tbl7.rows[ri].cells[si + 1].text)
        if val is not None and matched == 'PAHs':
            all_data.append([paper_id, site, 'Sum_PAH_ngg', val, 'ng/g', 'Table SI-7', 'soil', 'other', 'Tibet',
                           f'15 EPA PAH monomer sum from Table SI-7'])

# Table SI-8 (index 8): Soil zone V
tbl8 = doc.tables[8]
si8_sites = []
for c in tbl8.rows[1].cells[1:]:
    n = c.text.strip()
    si8_sites.append(n)
# Fix: 4th site has empty name - this is "Lhasa" (based on paper context)
# The paper mentions Lhasa, Xigaze, Nyemo, Lhaze, Mt.Everest, Saga, Zhongba, Wure, Nam co, Qurong
# The order from SI table: Xigaze, Nyemo, Lhaze, [empty], Qurong, Mt.Everest, Saga, Zhongba, Wure, Nam co
# The empty cell is Lhasa
for i in range(len(si8_sites)):
    if not si8_sites[i]:
        si8_sites[i] = 'Lhasa'

print(f'SI-8 sites (fixed): {si8_sites}')
si8_individual = [s for s in si8_sites if s and s.lower() != 'average']
print(f'SI-8 individual sites: {si8_individual}')

for ri in range(2, 18):
    comp = tbl8.rows[ri].cells[0].text.strip()
    if not comp:
        row_idx = ri - 2
        if row_idx < len(compounds):
            comp = compounds[row_idx]
        else:
            continue
    matched = None
    for c in compounds:
        if c in comp or c.lower() in comp.lower():
            matched = c
            break
    if not matched:
        continue
    for si, site in enumerate(si8_sites):
        if site.lower() == 'average':
            continue
        if si + 1 >= len(tbl8.rows[ri].cells):
            break
        val = parse_value(tbl8.rows[ri].cells[si + 1].text)
        if val is not None and matched == 'PAHs':
            all_data.append([paper_id, site, 'Sum_PAH_ngg', val, 'ng/g', 'Table SI-8', 'soil', 'other', 'Tibet',
                           f'15 EPA PAH monomer sum from Table SI-8'])

# Table SI-9 (index 9): Soil zones VI, VII, VIII
tbl9 = doc.tables[9]
si9_sites = []
for c in tbl9.rows[1].cells[1:]:
    n = c.text.strip()
    si9_sites.append(n)
print(f'SI-9 sites: {si9_sites}')
si9_individual = [s for s in si9_sites if s and s.lower() != 'average']
print(f'SI-9 individual sites: {si9_individual}')

for ri in range(2, 18):
    comp = tbl9.rows[ri].cells[0].text.strip()
    if not comp:
        row_idx = ri - 2
        if row_idx < len(compounds):
            comp = compounds[row_idx]
        else:
            continue
    matched = None
    for c in compounds:
        if c in comp or c.lower() in comp.lower():
            matched = c
            break
    if not matched:
        continue
    for si, site in enumerate(si9_sites):
        if site.lower() == 'average':
            continue
        if si + 1 >= len(tbl9.rows[ri].cells):
            break
        val = parse_value(tbl9.rows[ri].cells[si + 1].text)
        if val is not None and matched == 'PAHs':
            all_data.append([paper_id, site, 'Sum_PAH_ngg', val, 'ng/g', 'Table SI-9', 'soil', 'other', 'Tibet',
                           f'15 EPA PAH monomer sum from Table SI-9'])

# Print summary
print(f'\nTotal rows: {len(all_data)}')

# Verify: count per table
from collections import Counter
table_counts = Counter()
for row in all_data:
    tbl = row[5]
    table_counts[tbl] += 1
for k, v in sorted(table_counts.items()):
    print(f'  {k}: {v} rows')

# Print all data
for row in all_data:
    print(f'  {row[1]}: Sum_PAH={row[3]} ng/g [{row[5]}]')

# Check for any non-numeric values
for row in all_data:
    assert isinstance(row[3], (int, float)), f'Invalid value: {row}'

# Write CSV
out_dir = r'C:\Users\曾鸿\Desktop\SRS\outputs\literature_mining\manual_extract\op_only'
os.makedirs(out_dir, exist_ok=True)
out_file = os.path.join(out_dir, 'P07060.csv')

with open(out_file, 'w', newline='', encoding='utf-8-sig') as f:
    writer = csv.writer(f)
    writer.writerow(['paper_id', 'sample_id', 'pollutant_std', 'value', 'unit', 'evidence_location', 'matrix', 'site_type', 'province', 'extract_notes'])
    writer.writerows(all_data)

# Verify CSV can be read back
with open(out_file, 'r', encoding='utf-8-sig') as f:
    reader = csv.reader(f)
    header = next(reader)
    rows = list(reader)
    print(f'\nCSV verification: {len(rows)} data rows, header={header}')

print(f'\nOutput: {out_file}')
print(f'Done. {len(all_data)} data rows extracted.')
