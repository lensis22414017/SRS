"""从两份人工确认结构的DOCX SI提取可信样点级污染物数据。"""
from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document


OUT_DIR = Path(__file__).resolve().parents[1] / "trusted_reextract"
COAL_PATH = Path(
    r"G:\所有文献\8.第八阶段挖掘\si\DataSheet1_Characterization of polycyclic "
    r"aromatic hydrocarbons in soil in a coal mining area, East China_ Spatial "
    r"distribution, sources, and carcinog.docx"
)
INDUSTRIAL_PATH = OUT_DIR / "industrial_hm_pah_si" / "1-s2.0-S2666498422000254-mmc3.docx"
HM_ANALYTES = {"Cd", "Pb", "As", "Cr", "Hg", "Cu", "Zn", "Ni"}
PAH_ANALYTES = {
    "NAP", "NaP", "ACY", "Acy", "Ace", "FLO", "PHE", "ANT", "FLA", "PYR",
    "BaA", "CHR", "BbF", "BkF", "BaP", "DhA", "BgP", "IcP", "InP", "SumPAH",
}


def numeric_value(text: str) -> tuple[float, bool]:
    clean = str(text).strip()
    if not clean or re.fullmatch(r"(?i)n\.?\s*d\.?#?", clean):
        return 0.0, True
    match = re.search(r"[-+]?\d+(?:,\d{3})*(?:\.\d+)?", clean)
    if not match:
        raise ValueError(f"无法解析浓度值: {text!r}")
    return float(match.group().replace(",", "")), False


def collapse_consecutive(values: list[str]) -> list[tuple[int, str]]:
    result: list[tuple[int, str]] = []
    previous = object()
    for index, value in enumerate(values):
        value = value.replace("\n", " ").strip()
        if value != previous:
            result.append((index, value))
            previous = value
    return result


def normalize_analyte(value: str) -> str:
    name = re.sub(r"\s*\([^)]*\)\s*", "", value).strip()
    if name in {"ΣPAHs", "∑PAHs", "PAHs"}:
        return "SumPAH"
    return name


def extract_coal_pah(path: Path = COAL_PATH) -> pd.DataFrame:
    document = Document(path)
    records: list[dict] = []
    for table_index in (0, 1):
        table = document.tables[table_index]
        header = [cell.text.strip() for cell in table.rows[0].cells]
        sample_columns = [(index, value) for index, value in enumerate(header) if re.fullmatch(r"(?:US|CS)\d+", value)]
        for row_index, row in enumerate(table.rows[2:], start=2):
            cells = [cell.text.strip() for cell in row.cells]
            analyte = normalize_analyte(cells[0])
            if not analyte or analyte.lower() in {"mean", "sd"}:
                continue
            for column_index, sample_id in sample_columns:
                value, below = numeric_value(cells[column_index])
                records.append({
                    "source_id": "coal_mining_east_china_pah_si",
                    "sample_id": sample_id, "matrix": "soil",
                    "pollutant_family": "PAH", "analyte": analyte,
                    "value": value, "unit": "ng/g", "below_detection": below,
                    "evidence_location": f"DOCX:Table{table_index}:row{row_index}:col{column_index}",
                    "source_path": str(path), "evidence_level": "A_sample_table",
                })
    return pd.DataFrame(records).drop_duplicates(["source_id", "sample_id", "analyte"])


def extract_industrial_hm_pah(path: Path = INDUSTRIAL_PATH) -> pd.DataFrame:
    document = Document(path)
    records: list[dict] = []
    for table_index in (0, 1):
        active_headers: list[tuple[int, str]] | None = None
        for row_index, row in enumerate(document.tables[table_index].rows):
            raw = [cell.text.strip() for cell in row.cells]
            collapsed = collapse_consecutive(raw)
            first = collapsed[0][1] if collapsed else ""
            if first == "Sample":
                active_headers = [(index, normalize_analyte(value)) for index, value in collapsed]
                continue
            if active_headers is None or first not in {"A", "B", "C", "D", "E"}:
                continue
            for column_index, analyte in active_headers[1:]:
                if analyte not in HM_ANALYTES | PAH_ANALYTES:
                    continue
                if analyte == "SumPAH":
                    text = next((value for value in reversed(raw) if value.strip()), "")
                elif column_index < len(raw):
                    text = raw[column_index]
                else:
                    continue
                value, below = numeric_value(text)
                family = "HM" if analyte in HM_ANALYTES else "PAH"
                records.append({
                    "source_id": "industrial_sites_hm_pah_si",
                    "sample_id": first, "matrix": "soil",
                    "pollutant_family": family, "analyte": analyte,
                    "value": value, "unit": "mg/kg", "below_detection": below,
                    "evidence_location": f"DOCX:Table{table_index}:row{row_index}",
                    "source_path": str(path), "evidence_level": "A_sample_table",
                })
    data = pd.DataFrame(records)
    return data.drop_duplicates(["source_id", "sample_id", "analyte"], keep="first")


def write_outputs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    coal = extract_coal_pah()
    industrial = extract_industrial_hm_pah()
    coal.to_csv(OUT_DIR / "coal_mining_pah_trusted_long.csv", index=False, encoding="utf-8-sig")
    industrial.to_csv(OUT_DIR / "industrial_hm_pah_trusted_long.csv", index=False, encoding="utf-8-sig")
    combined = pd.concat([coal, industrial], ignore_index=True)
    combined.to_csv(OUT_DIR / "trusted_docx_si_combined_long.csv", index=False, encoding="utf-8-sig")
    wide = combined.pivot_table(
        index=["source_id", "sample_id", "matrix"], columns="analyte", values="value", aggfunc="first"
    ).reset_index()
    wide.to_csv(OUT_DIR / "trusted_docx_si_combined_wide.csv", index=False, encoding="utf-8-sig")
    print(f"coal_rows={len(coal)} coal_samples={coal.sample_id.nunique()}")
    print(f"industrial_rows={len(industrial)} industrial_samples={industrial.sample_id.nunique()}")


if __name__ == "__main__":
    write_outputs()
